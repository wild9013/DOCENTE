import pandas as pd
from docx import Document
from docx.shared import Pt

# Cargar los datos de los estudiantes y respuestas correctas
ruta_estudiantes = 'estudiantes.xlsx'
ruta_respuestas_correctas = 'respuestas_correctas.xlsx'

df_estudiantes = pd.read_excel(ruta_estudiantes)
df_respuestas_correctas = pd.read_excel(ruta_respuestas_correctas)

# Crear un set para almacenar todas las áreas
todas_las_areas = set(df_estudiantes['area'])

# Convertir set en lista ordenada
lista_areas = sorted(list(todas_las_areas))

# Agrupar por estudiante
estudiantes_grupo = df_estudiantes.groupby('nombre estudiante')

# Recorrer cada estudiante y generar un documento individual
for nombre, grupo_estudiante in estudiantes_grupo:
    # Crear un documento de Word para cada estudiante
    documento = Document()
    
    # Título del documento
    documento.add_heading(f'Informe de Notas - {nombre}', level=1)
    
    # Crear la primera fila de la tabla con los encabezados dinámicos (Áreas)
    encabezados = ['Área', 'Nota']
    tabla = documento.add_table(rows=1, cols=2)
    tabla.style = 'Table Grid'
    
    # Llenar los encabezados de la tabla
    encabezado_fila = tabla.rows[0].cells
    encabezado_fila[0].text = 'Área'
    encabezado_fila[1].text = 'Nota'
    
    # Crear un diccionario para almacenar las notas por área del estudiante
    notas_por_area = {area: '' for area in lista_areas}
    
    # Recorrer cada área en la que el estudiante tiene exámenes
    for index, fila_estudiante in grupo_estudiante.iterrows():
        area = fila_estudiante['area']
        respuestas_correctas = df_respuestas_correctas[df_respuestas_correctas['area'] == area].iloc[0]

        # Contador de respuestas correctas
        respuestas_correctas_contador = 0

        # Comparar respuestas
        for i in range(1, 11):
            if fila_estudiante[f'pregunta {i}'] == respuestas_correctas[f'pregunta {i}']:
                respuestas_correctas_contador += 1
        
        # Calcular la nota final sobre 5
        definitiva = (respuestas_correctas_contador / 10) * 5
        # Guardar la nota en el diccionario para el área correspondiente
        notas_por_area[area] = f'{definitiva:.2f}'
    
    # Añadir las notas por área a la tabla
    for area, nota in notas_por_area.items():
        fila = tabla.add_row().cells
        fila[0].text = area
        fila[1].text = nota

    # Ajustar el tamaño de fuente para toda la tabla
    for fila in tabla.rows:
        for celda in fila.cells:
            for parrafo in celda.paragraphs:
                for run in parrafo.runs:
                    run.font.size = Pt(12)

    # Guardar el documento de Word para el estudiante
    documento.save(f'informe_{nombre.replace(" ", "_")}.docx')

print("Informes generados exitosamente.")
