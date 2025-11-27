from docx import Document
from docx.shared import Pt

# Crear un documento de Word
documento = Document()

# Título del documento
documento.add_heading('Informe de Notas por Área', level=1)

# Lista de estudiantes y áreas de ejemplo
estudiantes = [
    ('Juan Pérez', 'Matemáticas'),
    ('María Gómez', 'Ciencias'),
    ('Carlos Ramírez', 'Matemáticas')
]

# Añadir la tabla
tabla = documento.add_table(rows=1, cols=3)
tabla.style = 'Table Grid'

# Crear encabezados de tabla
encabezados = tabla.rows[0].cells
encabezados[0].text = 'Nombre del Estudiante'
encabezados[1].text = 'Área'
encabezados[2].text = 'Nota Definitiva'

# Añadir los datos de los estudiantes
for estudiante in estudiantes:
    fila = tabla.add_row().cells
    fila[0].text = estudiante[0]
    fila[1].text = estudiante[1]
    fila[2].text = ''  # Espacio para llenar la nota definitiva

# Añadir un estilo de fuente más grande a la tabla
for fila in tabla.rows:
    for celda in fila.cells:
        for parrafo in celda.paragraphs:
            run = parrafo.runs
            for r in run:
                r.font.size = Pt(12)

# Guardar el documento
documento.save('plantilla_notas.docx')

print("Plantilla en Word creada exitosamente.")
