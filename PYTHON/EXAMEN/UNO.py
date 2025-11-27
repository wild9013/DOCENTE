import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
import os
from docx.shared import Cm, Pt
import docx

def generar_informes(excel_path, plantilla_estudiante_path, plantilla_grado_path, output_dir):
    os.makedirs(output_dir, exist_ok=True)
    
    # Leer y procesar datos
    df = pd.read_excel(excel_path)
    correctas = obtener_respuestas_correctas(df)
    estudiantes_df = procesar_estudiantes(df, correctas)
    
    # Generar informes
    generar_informes_estudiantes(estudiantes_df, correctas, plantilla_estudiante_path, output_dir)
    generar_informe_grado(estudiantes_df, correctas, plantilla_grado_path, output_dir)

def obtener_respuestas_correctas(df):
    return df.iloc[0, 1:21].tolist()

def procesar_estudiantes(df, correctas):
    estudiantes_df = df.iloc[1:].dropna(subset=['nombre'])
    estudiantes_df['Correctas'] = estudiantes_df.apply(
        lambda row: sum(row[f'p{i}'] == correctas[i-1] for i in range(1, 21)),
        axis=1
    )
    estudiantes_df['Porcentaje'] = (estudiantes_df['Correctas'] / 20 * 100).round(1)
    return estudiantes_df

def generar_informes_estudiantes(estudiantes_df, correctas, plantilla_path, output_dir):
    for _, estudiante in estudiantes_df.iterrows():
        doc = Document(plantilla_path)
        
        # Reemplazar placeholders
        reemplazar_placeholders(doc, {
            '{nombre}': estudiante['nombre'],
            '{total_correctas}': str(estudiante['Correctas']),
            '{porcentaje}': f"{estudiante['Porcentaje']}%",
            '{total_preguntas}': '20'
        })
        
        # Generar contenido
        generar_grafico_resumen(estudiante, output_dir)
        generar_tabla_respuestas(doc, estudiante, correctas)
        doc.add_picture(os.path.join(output_dir, f'resumen_{estudiante["nombre"]}.png'), width=Cm(12))
        
        # Guardar y limpiar
        nombre_archivo = estudiante['nombre'].replace(' ', '_')
        doc.save(os.path.join(output_dir, f'Informe_{nombre_archivo}.docx'))
        os.remove(os.path.join(output_dir, f'resumen_{estudiante["nombre"]}.png'))

def generar_grafico_resumen(estudiante, output_dir):
    plt.figure(figsize=(8, 4))
    categorias = ['Correctas', 'Incorrectas']
    valores = [estudiante['Correctas'], 20 - estudiante['Correctas']]
    
    plt.bar(categorias, valores, color=['#4CAF50', '#F44336'])
    plt.title('Resumen de respuestas')
    plt.ylabel('Cantidad')
    plt.ylim(0, 20)
    
    for i, v in enumerate(valores):
        plt.text(i, v + 0.5, str(v), ha='center')
    
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, f'resumen_{estudiante["nombre"]}.png'), bbox_inches='tight')
    plt.close()

def generar_tabla_respuestas(doc, estudiante, correctas):
    # Crear tabla básica sin estilo específico
    tabla = doc.add_table(rows=1, cols=4)
    
    # Configurar encabezados
    encabezados = tabla.rows[0].cells
    encabezados[0].text = 'Pregunta'
    encabezados[1].text = 'Respuesta'
    encabezados[2].text = 'Correcta'
    encabezados[3].text = 'Estado'

    # Añadir datos
    for i in range(1, 21):
        respuesta = estudiante[f'p{i}']
        correcta = correctas[i-1]
        estado = '✅' if respuesta == correcta else '❌'
        
        fila = tabla.add_row().cells
        fila[0].text = str(i)
        fila[1].text = str(respuesta)
        fila[2].text = str(correcta)
        fila[3].text = estado

    # Ajustar formato
    for row in tabla.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Ajustar ancho columnas
    for col in tabla.columns:
        col.width = Cm(3.5)

def generar_informe_grado(estudiantes_df, correctas, plantilla_path, output_dir):
    doc = Document(plantilla_path)
    
    # Estadísticas
    stats = {
        '{promedio_grado}': f"{estudiantes_df['Porcentaje'].mean():.1f}%",
        '{maxima_nota}': f"{estudiantes_df['Porcentaje'].max():.1f}%",
        '{minima_nota}': f"{estudiantes_df['Porcentaje'].min():.1f}%",
        '{tasa_aprobacion}': f"{(estudiantes_df['Porcentaje'] >= 60).mean() * 100:.1f}%"
    }
    reemplazar_placeholders(doc, stats)
    
    # Gráficos
    generar_grafico_distribucion(estudiantes_df, output_dir)
    generar_grafico_dificultad(estudiantes_df, correctas, output_dir)
    
    doc.add_paragraph("\nDistribución de resultados:")
    doc.add_picture(os.path.join(output_dir, 'distribucion.png'), width=Cm(12))
    
    doc.add_paragraph("\nDificultad por pregunta:")
    doc.add_picture(os.path.join(output_dir, 'dificultad.png'), width=Cm(12))
    
    # Guardar
    doc.save(os.path.join(output_dir, 'Informe_Grado.docx'))
    for archivo in ['distribucion.png', 'dificultad.png']:
        os.remove(os.path.join(output_dir, archivo))

def generar_grafico_distribucion(estudiantes_df, output_dir):
    plt.figure(figsize=(10, 5))
    plt.hist(estudiantes_df['Porcentaje'], bins=10, edgecolor='black', color='#2196F3')
    plt.title('Distribución de resultados')
    plt.xlabel('Porcentaje')
    plt.ylabel('Estudiantes')
    plt.grid(axis='y', alpha=0.75)
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, 'distribucion.png'))
    plt.close()

def generar_grafico_dificultad(estudiantes_df, correctas, output_dir):
    dificultad = []
    for i in range(1, 21):
        col = f'p{i}'
        correctos = (estudiantes_df[col] == correctas[i-1]).sum()
        dificultad.append((correctos / len(estudiantes_df)) * 100)
    
    plt.figure(figsize=(12, 6))
    plt.bar(range(1, 21), dificultad, color='#FF9800')
    plt.title('Dificultad por pregunta')
    plt.xlabel('Pregunta')
    plt.ylabel('% Correctas')
    plt.ylim(0, 100)
    plt.xticks(range(1, 21))
    plt.grid(axis='y', alpha=0.75)
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, 'dificultad.png'))
    plt.close()

def reemplazar_placeholders(doc, replacements):
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, value)
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

# Configuración
if __name__ == "__main__":
    excel_path = 'respuestas.xlsx'
    plantilla_estudiante = 'plantilla_estudiante.docx'
    plantilla_grado = 'plantilla_grado.docx'
    output_directory = 'informes_generados'
    
    generar_informes(excel_path, plantilla_estudiante, plantilla_grado, output_directory)