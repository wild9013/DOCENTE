import tkinter as tk
from tkinter import messagebox, filedialog, ttk
import random
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime, timedelta
import os

# --- Configuración de Estilo Global ---
class Theme:
    BG_COLOR = "#f0f2f5"
    PRIMARY_COLOR = "#2c3e50"
    ACCENT_COLOR = "#3498db"
    CORRECT_COLOR = "#27ae60"
    INCORRECT_COLOR = "#e74c3c"
    TEXT_COLOR = "#34495e"
    WHITE = "#ffffff"
    
    FONT_TITLE = ("Segoe UI", 24, "bold")
    FONT_SUBTITLE = ("Segoe UI", 16)
    FONT_NORMAL = ("Segoe UI", 12)
    FONT_BOLD = ("Segoe UI", 12, "bold")
    FONT_BIG_BOLD = ("Segoe UI", 14, "bold")

class SistemaEvaluacion:
    def __init__(self, root):
        self.root = root
        self.root.title("Sistema de Evaluación - Trivia Educativa")
        self.root.geometry("900x650")
        self.root.resizable(False, False)
        self.root.configure(bg=Theme.BG_COLOR)
        
        # Centrar ventana
        self.center_window(900, 650)
        
        # Variables de estado
        self.file_path = None
        self.cover_image_path = None
        self.data = None
        self.student_name = ""
        self.current_question = 0
        self.correct_answers = 0
        self.responses = []
        self.start_time = None
        self.timer_running = False
        
        # Iniciar
        self.show_file_selection()
    
    def center_window(self, width, height):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width // 2) - (width // 2)
        y = (screen_height // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')

    def clear_window(self):
        """Limpia todos los widgets de la ventana actual"""
        for widget in self.root.winfo_children():
            widget.destroy()

    # --- PANTALLA 1: CONFIGURACIÓN ---
    def show_file_selection(self):
        self.clear_window()
        
        # Contenedor principal centrado
        main_frame = tk.Frame(self.root, bg=Theme.BG_COLOR)
        main_frame.pack(expand=True, fill="both", padx=50, pady=40)

        # Header
        tk.Label(main_frame, text="Configuración del Examen", font=Theme.FONT_TITLE, 
                 bg=Theme.BG_COLOR, fg=Theme.PRIMARY_COLOR).pack(pady=(0, 10))
        
        tk.Label(main_frame, text="Sube tu banco de preguntas para comenzar", font=Theme.FONT_SUBTITLE, 
                 bg=Theme.BG_COLOR, fg="#7f8c8d").pack(pady=(0, 30))

        # Tarjeta blanca para controles
        card = tk.Frame(main_frame, bg=Theme.WHITE, bd=1, relief="solid")
        card.pack(fill="x", ipady=20, padx=20)

        # Sección Excel
        tk.Label(card, text="1. Banco de Preguntas (Excel)", font=Theme.FONT_BOLD, 
                 bg=Theme.WHITE, fg=Theme.PRIMARY_COLOR).pack(pady=(10, 5))
        
        self.excel_status = tk.Label(card, text="No seleccionado", font=("Segoe UI", 10), 
                                   bg=Theme.WHITE, fg="#95a5a6")
        self.excel_status.pack(pady=5)
        
        tk.Button(card, text="📂 Cargar Excel", command=self.select_excel_file,
                  font=Theme.FONT_NORMAL, bg=Theme.ACCENT_COLOR, fg=Theme.WHITE,
                  relief="flat", cursor="hand2", width=20).pack(pady=5)

        tk.Frame(card, height=1, bg="#ecf0f1").pack(fill="x", pady=20, padx=40) # Separador

        # Sección Imagen
        tk.Label(card, text="2. Portada del Reporte (Opcional)", font=Theme.FONT_BOLD, 
                 bg=Theme.WHITE, fg=Theme.PRIMARY_COLOR).pack(pady=(0, 5))
        
        self.image_status = tk.Label(card, text="Opcional", font=("Segoe UI", 10), 
                                   bg=Theme.WHITE, fg="#95a5a6")
        self.image_status.pack(pady=5)
        
        tk.Button(card, text="🖼️ Cargar Imagen", command=self.select_cover_image,
                  font=Theme.FONT_NORMAL, bg="#95a5a6", fg=Theme.WHITE,
                  relief="flat", cursor="hand2", width=20).pack(pady=5)

        # Botón Continuar
        self.continue_btn = tk.Button(main_frame, text="Continuar ➜", command=self.show_student_form,
                                      font=Theme.FONT_BIG_BOLD, bg=Theme.CORRECT_COLOR, fg=Theme.WHITE,
                                      state="disabled", cursor="arrow", relief="flat", pady=10)
        self.continue_btn.pack(pady=30, fill="x", padx=100)

    def select_excel_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx *.xls")])
        if file_path:
            self.file_path = file_path
            if self.load_questions():
                name = os.path.basename(file_path)
                self.excel_status.config(text=f"✓ {name} ({len(self.data)} preguntas)", fg=Theme.CORRECT_COLOR)
                self.continue_btn.config(state="normal", cursor="hand2")
            else:
                self.excel_status.config(text="⚠ Error en el archivo", fg=Theme.INCORRECT_COLOR)

    def select_cover_image(self):
        file_path = filedialog.askopenfilename(filetypes=[("Imágenes", "*.png *.jpg *.jpeg")])
        if file_path:
            self.cover_image_path = file_path
            self.image_status.config(text=f"✓ {os.path.basename(file_path)}", fg=Theme.CORRECT_COLOR)

    def load_questions(self):
        try:
            df = pd.read_excel(self.file_path)
            # Normalizar columnas (eliminar espacios extra)
            df.columns = df.columns.str.strip()
            
            required = ["Pregunta", "Respuesta Correcta", "R1", "R2", "R3"]
            if not all(col in df.columns for col in required):
                messagebox.showerror("Error de Formato", f"Faltan columnas.\nRequeridas: {required}")
                return False
            
            # Limpieza y Aleatorización de preguntas
            self.data = df.dropna(subset=["Pregunta"]).astype(str).sample(frac=1).reset_index(drop=True)
            return len(self.data) > 0
        except Exception as e:
            messagebox.showerror("Error", str(e))
            return False

    # --- PANTALLA 2: DATOS ESTUDIANTE ---
    def show_student_form(self):
        self.clear_window()
        frame = tk.Frame(self.root, bg=Theme.BG_COLOR)
        frame.pack(expand=True, fill="both", padx=100, pady=80)

        tk.Label(frame, text="Datos del Estudiante", font=Theme.FONT_TITLE, 
                 bg=Theme.BG_COLOR, fg=Theme.PRIMARY_COLOR).pack(pady=20)

        entry_frame = tk.Frame(frame, bg=Theme.WHITE, padx=20, pady=20, relief="solid", bd=1)
        entry_frame.pack(fill="x")

        tk.Label(entry_frame, text="Nombre Completo:", font=Theme.FONT_BOLD, 
                 bg=Theme.WHITE, fg=Theme.TEXT_COLOR).pack(anchor="w")
        
        self.name_entry = tk.Entry(entry_frame, font=("Segoe UI", 14), bg="#f9f9f9", relief="solid", bd=1)
        self.name_entry.pack(fill="x", pady=10, ipady=5)
        self.name_entry.focus()
        self.name_entry.bind("<Return>", lambda e: self.start_exam())

        btn_frame = tk.Frame(frame, bg=Theme.BG_COLOR)
        btn_frame.pack(pady=30, fill="x")

        tk.Button(btn_frame, text="← Volver", command=self.show_file_selection,
                  font=Theme.FONT_NORMAL, bg="#7f8c8d", fg=Theme.WHITE, relief="flat", width=15).pack(side="left")
        
        tk.Button(btn_frame, text="Comenzar Evaluación", command=self.start_exam,
                  font=Theme.FONT_BOLD, bg=Theme.PRIMARY_COLOR, fg=Theme.WHITE, relief="flat", width=20).pack(side="right")

    # --- PANTALLA 3: EXAMEN (CORE) ---
    def start_exam(self):
        name = self.name_entry.get().strip()
        if not name:
            messagebox.showwarning("Atención", "Por favor escribe tu nombre.")
            return
        
        self.student_name = name
        self.start_time = datetime.now()
        self.current_question = 0
        self.correct_answers = 0
        self.responses = []
        self.timer_running = True
        self.show_question()
        self.update_timer()

    def update_timer(self):
        if self.timer_running:
            elapsed = datetime.now() - self.start_time
            # Formato MM:SS
            minutes, seconds = divmod(elapsed.seconds, 60)
            time_str = f"{minutes:02}:{seconds:02}"
            
            try:
                self.timer_label.config(text=f"⏱ {time_str}")
                self.root.after(1000, self.update_timer)
            except AttributeError:
                pass # Si la etiqueta no existe (cambio de pantalla), ignorar

    def show_question(self):
        if self.current_question >= len(self.data):
            self.show_results()
            return

        self.clear_window()
        
        # --- Header (Progreso y Timer) ---
        header = tk.Frame(self.root, bg=Theme.PRIMARY_COLOR, height=60, padx=20)
        header.pack(fill="x")
        header.pack_propagate(False)

        prog_text = f"Pregunta {self.current_question + 1} / {len(self.data)}"
        tk.Label(header, text=prog_text, font=Theme.FONT_BOLD, bg=Theme.PRIMARY_COLOR, fg=Theme.WHITE).pack(side="left")
        
        self.timer_label = tk.Label(header, text="⏱ 00:00", font=Theme.FONT_BOLD, bg=Theme.PRIMARY_COLOR, fg="#f1c40f")
        self.timer_label.pack(side="right")

        # Barra de progreso visual
        progress_frame = tk.Frame(self.root, height=5, bg="#ecf0f1")
        progress_frame.pack(fill="x")
        pct = (self.current_question / len(self.data))
        tk.Frame(progress_frame, width=int(900 * pct), height=5, bg=Theme.ACCENT_COLOR).pack(side="left")

        # --- Área de Pregunta ---
        content = tk.Frame(self.root, bg=Theme.BG_COLOR)
        content.pack(expand=True, fill="both", padx=50, pady=20)

        q_card = tk.Frame(content, bg=Theme.WHITE, relief="solid", bd=0, padx=30, pady=30)
        q_card.pack(fill="x", pady=(0, 20))
        
        question_text = self.data.iloc[self.current_question]["Pregunta"]
        tk.Label(q_card, text=question_text, font=("Segoe UI", 16, "bold"), 
                 bg=Theme.WHITE, fg=Theme.TEXT_COLOR, wraplength=750, justify="left").pack(anchor="w")

        # --- Opciones ---
        self.options_frame = tk.Frame(content, bg=Theme.BG_COLOR)
        self.options_frame.pack(fill="x")

        row = self.data.iloc[self.current_question]
        options = [row["Respuesta Correcta"], row["R1"], row["R2"], row["R3"]]
        random.shuffle(options)
        
        self.option_buttons = []

        for idx, opt in enumerate(options):
            # Usamos una función lambda con valor por defecto para capturar el valor de 'opt' en el loop
            btn = tk.Button(self.options_frame, text=f"{chr(65+idx)}. {opt}", 
                           font=("Segoe UI", 13), bg=Theme.WHITE, fg=Theme.TEXT_COLOR,
                           anchor="w", padx=20, pady=12, relief="flat", bd=0,
                           activebackground="#ecf0f1", cursor="hand2",
                           command=lambda o=opt, b_idx=idx: self.handle_answer(o, b_idx))
            btn.pack(fill="x", pady=6)
            self.option_buttons.append(btn)

        # Botón saltar (discreto abajo)
        tk.Button(self.root, text="Saltar Pregunta", command=lambda: self.handle_answer(None, -1),
                  font=("Segoe UI", 10), bg=Theme.BG_COLOR, fg="#95a5a6", 
                  relief="flat", cursor="hand2").pack(side="bottom", pady=10)

    def handle_answer(self, selected_text, btn_index):
        # Desactivar todos los botones para evitar doble clic
        for btn in self.option_buttons:
            btn.config(state="disabled", cursor="arrow")

        correct_text = self.data.iloc[self.current_question]["Respuesta Correcta"]
        is_correct = (selected_text == correct_text)
        
        # Guardar respuesta
        self.responses.append({
            "pregunta": self.data.iloc[self.current_question]["Pregunta"],
            "respuesta_estudiante": selected_text if selected_text else "Saltada",
            "respuesta_correcta": correct_text,
            "correcta": is_correct
        })

        if is_correct:
            self.correct_answers += 1
            if btn_index >= 0:
                self.option_buttons[btn_index].config(bg=Theme.CORRECT_COLOR, fg=Theme.WHITE)
        else:
            if btn_index >= 0:
                self.option_buttons[btn_index].config(bg=Theme.INCORRECT_COLOR, fg=Theme.WHITE)
            
            # Mostrar cuál era la correcta
            for btn in self.option_buttons:
                if btn['text'].endswith(f" {correct_text}"): # Busca el texto (cuidado con el prefijo A. B. etc)
                    btn.config(bg=Theme.CORRECT_COLOR, fg=Theme.WHITE)

        # Esperar 1.5 segundos y pasar a la siguiente
        self.root.after(1500, self.next_step)

    def next_step(self):
        self.current_question += 1
        self.show_question()

    # --- PANTALLA 4: RESULTADOS ---
    def show_results(self):
        self.timer_running = False
        self.clear_window()
        
        elapsed = datetime.now() - self.start_time
        score_pct = (self.correct_answers / len(self.data)) * 100
        
        frame = tk.Frame(self.root, bg=Theme.BG_COLOR)
        frame.pack(expand=True, fill="both", padx=50, pady=40)

        # Título y Nombre
        tk.Label(frame, text="Evaluación Finalizada", font=Theme.FONT_TITLE, bg=Theme.BG_COLOR, fg=Theme.PRIMARY_COLOR).pack()
        tk.Label(frame, text=f"Estudiante: {self.student_name}", font=Theme.FONT_SUBTITLE, bg=Theme.BG_COLOR, fg="#7f8c8d").pack(pady=(0, 20))

        # Círculo de Puntuación (Simulado con Label)
        color = Theme.CORRECT_COLOR if score_pct >= 60 else Theme.INCORRECT_COLOR
        score_frame = tk.Frame(frame, bg=color, width=150, height=150) # Cuadrado redondeado visual
        # Truco para layout circular simple
        tk.Label(frame, text=f"{int(score_pct)}%", font=("Arial", 48, "bold"), fg=color, bg=Theme.BG_COLOR).pack(pady=10)
        
        msg = "¡Excelente Trabajo!" if score_pct >= 80 else ("Buen intento" if score_pct >= 60 else "Necesitas reforzar")
        tk.Label(frame, text=msg, font=Theme.FONT_BIG_BOLD, bg=Theme.BG_COLOR, fg=Theme.TEXT_COLOR).pack()

        # Grid de Estadísticas
        stats_frame = tk.Frame(frame, bg=Theme.BG_COLOR)
        stats_frame.pack(pady=30)
        
        self.create_stat_card(stats_frame, "Correctas", self.correct_answers, Theme.CORRECT_COLOR, 0)
        self.create_stat_card(stats_frame, "Incorrectas", len(self.data)-self.correct_answers, Theme.INCORRECT_COLOR, 1)
        minutes = elapsed.seconds // 60
        self.create_stat_card(stats_frame, "Tiempo", f"{minutes} min", Theme.ACCENT_COLOR, 2)

        # Botones finales
        btn_frame = tk.Frame(frame, bg=Theme.BG_COLOR)
        btn_frame.pack(pady=20)

        tk.Button(btn_frame, text="📄 Generar Reporte Word", command=self.generate_report,
                  font=Theme.FONT_BOLD, bg=Theme.PRIMARY_COLOR, fg=Theme.WHITE, 
                  padx=20, pady=10, relief="flat", cursor="hand2").pack(side="left", padx=10)
        
        tk.Button(btn_frame, text="Salir", command=self.root.quit,
                  font=Theme.FONT_NORMAL, bg="#95a5a6", fg=Theme.WHITE, 
                  padx=20, pady=10, relief="flat", cursor="hand2").pack(side="left", padx=10)

    def create_stat_card(self, parent, title, value, color, col_idx):
        f = tk.Frame(parent, bg=Theme.WHITE, padx=20, pady=15, relief="solid", bd=1)
        f.grid(row=0, column=col_idx, padx=10)
        tk.Label(f, text=str(value), font=("Arial", 20, "bold"), fg=color, bg=Theme.WHITE).pack()
        tk.Label(f, text=title, font=("Segoe UI", 10), fg="#7f8c8d", bg=Theme.WHITE).pack()

    # --- GENERACIÓN DE REPORTE ---
    def generate_report(self):
        try:
            doc = Document()
            # Estilos básicos
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            # Portada
            if self.cover_image_path:
                try:
                    doc.add_picture(self.cover_image_path, width=Inches(4))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except: pass

            title = doc.add_heading('REPORTE DE RESULTADOS', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Datos
            doc.add_heading('Información', level=1)
            p = doc.add_paragraph()
            p.add_run('Estudiante: ').bold = True
            p.add_run(f"{self.student_name}\n")
            p.add_run('Fecha: ').bold = True
            p.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M')}\n")
            p.add_run('Puntaje: ').bold = True
            score = (self.correct_answers / len(self.data)) * 100
            p.add_run(f"{score:.1f}% ({self.correct_answers}/{len(self.data)})")

            # Tabla de detalles
            doc.add_heading('Detalle de Preguntas', level=1)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Shading Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Pregunta'
            hdr_cells[1].text = 'Tu Respuesta'
            hdr_cells[2].text = 'Estado'

            for resp in self.responses:
                row_cells = table.add_row().cells
                row_cells[0].text = str(resp['pregunta'])
                row_cells[1].text = str(resp['respuesta_estudiante'])
                
                status = "Correcto" if resp['correcta'] else "Incorrecto"
                row_cells[2].text = status

            # --- CAMBIO PRINCIPAL AQUÍ ---
            # 1. Definir nombre sugerido
            default_name = f"Resultados_{self.student_name.replace(' ', '_')}_{datetime.now().strftime('%H%M')}"
            
            # 2. Abrir cuadro de diálogo para "Guardar como"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Documento Word", "*.docx")],
                initialfile=default_name,
                title="Guardar Reporte de Evaluación"
            )

            # 3. Si el usuario seleccionó una ruta (no dio cancelar)
            if file_path:
                doc.save(file_path)
                messagebox.showinfo("Éxito", f"Reporte guardado exitosamente en:\n{file_path}")
            else:
                # El usuario canceló la operación
                return
            
        except Exception as e:
            messagebox.showerror("Error al guardar", f"No se pudo guardar el archivo:\n{str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    # Intenta configurar icono si existe, si no, ignora
    # root.iconbitmap('icon.ico') 
    app = SistemaEvaluacion(root)
    root.mainloop()