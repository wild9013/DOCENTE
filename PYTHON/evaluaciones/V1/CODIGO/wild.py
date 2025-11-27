import tkinter as tk
from tkinter import messagebox, ttk, filedialog
import random
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime
import os

class TriviaGame:
    def __init__(self, root):
        self.root = root
        self.root.title("Juego Trivia Educativo")
        self.root.geometry("600x550")
        self.root.resizable(False, False)
        
        # Variables del juego
        self.question_index = 0
        self.correct_answers = 0
        self.answer_var = tk.StringVar()
        self.responses = []
        self.student_name = ""
        self.data = None
        self.shuffled_options = []
        self.file_path = None
        
        # Configurar estilos
        self.setup_styles()
        
        # Crear pantalla de selección de archivo
        self.create_file_selection_screen()
    
    def setup_styles(self):
        """Configurar estilos visuales"""
        style = ttk.Style()
        style.theme_use('clam')
        
        # Colores
        self.bg_color = "#f0f0f0"
        self.primary_color = "#4CAF50"
        self.error_color = "#f44336"
        self.text_color = "#333333"
        
        self.root.configure(bg=self.bg_color)
    
    def create_file_selection_screen(self):
        """Crear pantalla de selección de archivo"""
        self.file_frame = tk.Frame(self.root, bg=self.bg_color)
        self.file_frame.pack(expand=True, fill="both", padx=20, pady=20)
        
        # Título
        title = tk.Label(self.file_frame, text="🎓 Juego de Trivia", 
                        font=("Arial", 24, "bold"), bg=self.bg_color, fg=self.primary_color)
        title.pack(pady=(40, 20))
        
        # Instrucciones
        info = tk.Label(self.file_frame, 
                       text="Selecciona el archivo Excel con las preguntas",
                       font=("Arial", 12), bg=self.bg_color, fg=self.text_color)
        info.pack(pady=10)
        
        # Ruta del archivo seleccionado
        self.file_label = tk.Label(self.file_frame, text="Ningún archivo seleccionado", 
                                   font=("Arial", 10, "italic"), bg=self.bg_color, 
                                   fg="#666666", wraplength=500)
        self.file_label.pack(pady=20)
        
        # Botón para seleccionar archivo
        select_btn = tk.Button(self.file_frame, text="📁 Seleccionar Archivo Excel", 
                              command=self.select_file, font=("Arial", 13, "bold"),
                              bg="#2196F3", fg="white", padx=30, pady=12,
                              cursor="hand2", relief="raised", bd=0)
        select_btn.pack(pady=10)
        
        # Botón para continuar (inicialmente deshabilitado)
        self.continue_btn = tk.Button(self.file_frame, text="Continuar →", 
                                     command=self.proceed_to_name_screen,
                                     font=("Arial", 13, "bold"), bg=self.primary_color,
                                     fg="white", padx=30, pady=12, cursor="hand2",
                                     relief="raised", bd=0, state="disabled")
        self.continue_btn.pack(pady=10)
        
        # Información adicional
        help_text = tk.Label(self.file_frame, 
                           text="El archivo debe contener las columnas:\nPregunta, Respuesta Correcta, R1, R2, R3",
                           font=("Arial", 9), bg=self.bg_color, fg="#888888", justify="center")
        help_text.pack(pady=(30, 0))
    
    def select_file(self):
        """Abrir diálogo para seleccionar archivo"""
        file_path = filedialog.askopenfilename(
            title="Seleccionar archivo de preguntas",
            filetypes=[
                ("Archivos Excel", "*.xlsx *.xls"),
                ("Todos los archivos", "*.*")
            ],
            initialdir=os.getcwd()
        )
        
        if file_path:
            self.file_path = file_path
            # Mostrar nombre del archivo
            file_name = os.path.basename(file_path)
            self.file_label.config(text=f"✓ Archivo seleccionado: {file_name}", 
                                  fg=self.primary_color, font=("Arial", 10, "bold"))
            
            # Validar el archivo
            if self.load_data():
                self.continue_btn.config(state="normal")
                messagebox.showinfo("Éxito", 
                                  f"Archivo cargado correctamente.\n\n"
                                  f"Se encontraron {len(self.data)} preguntas.")
            else:
                self.file_path = None
                self.file_label.config(text="❌ Archivo inválido. Selecciona otro archivo.", 
                                      fg=self.error_color)
                self.continue_btn.config(state="disabled")
    
    def load_data(self):
        """Cargar datos desde Excel con manejo de errores"""
        try:
            if not self.file_path:
                return False
            
            if not os.path.exists(self.file_path):
                messagebox.showerror("Error", f"No se encontró el archivo '{self.file_path}'")
                return False
            
            self.data = pd.read_excel(self.file_path)
            self.data = self.data.astype(str)
            
            # Validar columnas requeridas
            required_columns = ["Pregunta", "Respuesta Correcta", "R1", "R2", "R3"]
            if not all(col in self.data.columns for col in required_columns):
                messagebox.showerror("Error", "El archivo Excel no tiene las columnas requeridas:\n" + 
                                   "\n".join(required_columns))
                return False
            
            # Eliminar filas vacías
            self.data = self.data.dropna(subset=["Pregunta"])
            
            if len(self.data) == 0:
                messagebox.showerror("Error", "El archivo no contiene preguntas válidas")
                return False
            
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Error al cargar el archivo: {str(e)}")
            return False
    
    def proceed_to_name_screen(self):
        """Proceder a la pantalla de ingreso de nombre"""
        self.file_frame.pack_forget()
        self.create_name_screen()
    
    def create_name_screen(self):
        """Crear pantalla de ingreso de nombre"""
        self.name_frame = tk.Frame(self.root, bg=self.bg_color)
        self.name_frame.pack(expand=True, fill="both", padx=20, pady=20)
        
        # Título
        title = tk.Label(self.name_frame, text="🎓 Juego de Trivia", 
                        font=("Arial", 24, "bold"), bg=self.bg_color, fg=self.primary_color)
        title.pack(pady=(40, 20))
        
        # Información
        info = tk.Label(self.name_frame, 
                       text=f"Prepárate para responder {len(self.data)} preguntas",
                       font=("Arial", 12), bg=self.bg_color, fg=self.text_color)
        info.pack(pady=10)
        
        # Campo de nombre
        name_label = tk.Label(self.name_frame, text="Ingresa tu nombre:", 
                             font=("Arial", 14), bg=self.bg_color, fg=self.text_color)
        name_label.pack(pady=(30, 10))
        
        self.name_entry = tk.Entry(self.name_frame, font=("Arial", 14), width=30, bd=2, relief="solid")
        self.name_entry.pack(pady=10, ipady=5)
        self.name_entry.focus()
        self.name_entry.bind("<Return>", lambda e: self.start_game())
        
        # Botón comenzar
        start_btn = tk.Button(self.name_frame, text="Comenzar Juego", 
                            command=self.start_game, font=("Arial", 14, "bold"),
                            bg=self.primary_color, fg="white", padx=30, pady=10,
                            cursor="hand2", relief="raised", bd=0)
        start_btn.pack(pady=30)
    
    def create_question_screen(self):
        """Crear pantalla de preguntas"""
        self.question_frame = tk.Frame(self.root, bg=self.bg_color)
        self.question_frame.pack(expand=True, fill="both", padx=20, pady=20)
        
        # Encabezado con progreso
        header_frame = tk.Frame(self.question_frame, bg=self.bg_color)
        header_frame.pack(fill="x", pady=(0, 20))
        
        self.status_label = tk.Label(header_frame, text="", font=("Arial", 12, "bold"),
                                     bg=self.bg_color, fg=self.primary_color)
        self.status_label.pack(side="left")
        
        self.score_label = tk.Label(header_frame, text="Correctas: 0", 
                                    font=("Arial", 12, "bold"),
                                    bg=self.bg_color, fg=self.primary_color)
        self.score_label.pack(side="right")
        
        # Barra de progreso
        self.progress = ttk.Progressbar(self.question_frame, length=560, mode='determinate')
        self.progress.pack(pady=(0, 20))
        
        # Pregunta
        question_container = tk.Frame(self.question_frame, bg="white", relief="solid", bd=1)
        question_container.pack(fill="x", pady=(0, 20))
        
        self.question_label = tk.Label(question_container, text="", wraplength=540,
                                       font=("Arial", 14, "bold"), bg="white", 
                                       fg=self.text_color, justify="left", pady=20, padx=20)
        self.question_label.pack()
        
        # Opciones
        self.radio_buttons = []
        for i in range(4):
            rb_frame = tk.Frame(self.question_frame, bg="white", relief="solid", bd=1)
            rb_frame.pack(fill="x", pady=5)
            
            rb = tk.Radiobutton(rb_frame, text="", variable=self.answer_var, 
                               wraplength=500, font=("Arial", 12), bg="white",
                               fg=self.text_color, anchor="w", justify="left",
                               padx=15, pady=10, cursor="hand2")
            rb.pack(fill="x")
            self.radio_buttons.append(rb)
        
        # Botones
        button_frame = tk.Frame(self.question_frame, bg=self.bg_color)
        button_frame.pack(pady=20)
        
        self.answer_button = tk.Button(button_frame, text="✓ Responder", 
                                       command=self.handle_answer,
                                       font=("Arial", 13, "bold"), bg=self.primary_color,
                                       fg="white", padx=25, pady=10, cursor="hand2",
                                       relief="raised", bd=0)
        self.answer_button.pack(side="left", padx=5)
        
        skip_button = tk.Button(button_frame, text="⊗ Saltar", 
                               command=self.skip_question,
                               font=("Arial", 13, "bold"), bg="#ff9800",
                               fg="white", padx=25, pady=10, cursor="hand2",
                               relief="raised", bd=0)
        skip_button.pack(side="left", padx=5)
    
    def start_game(self):
        """Iniciar el juego"""
        self.student_name = self.name_entry.get().strip()
        if not self.student_name:
            messagebox.showwarning("Nombre Requerido", 
                                  "Por favor, ingresa tu nombre antes de comenzar.")
            self.name_entry.focus()
            return
        
        self.name_frame.pack_forget()
        self.create_question_screen()
        self.show_question()
    
    def show_question(self):
        """Mostrar pregunta actual"""
        if self.question_index < len(self.data):
            # Mezclar opciones
            correct = self.data["Respuesta Correcta"][self.question_index]
            options = [correct, 
                      self.data["R1"][self.question_index],
                      self.data["R2"][self.question_index],
                      self.data["R3"][self.question_index]]
            random.shuffle(options)
            self.shuffled_options = options
            
            # Actualizar interfaz
            self.question_label.config(text=self.data["Pregunta"][self.question_index])
            
            for i, option in enumerate(options):
                self.radio_buttons[i].config(text=f"{chr(65+i)}. {option}", value=option)
            
            self.answer_var.set("")  # Limpiar selección
            self.update_status()
            self.update_progress()
        else:
            self.end_game()
    
    def update_status(self):
        """Actualizar estado de la pregunta"""
        self.status_label.config(text=f"Pregunta {self.question_index + 1} de {len(self.data)}")
        self.score_label.config(text=f"Correctas: {self.correct_answers}")
    
    def update_progress(self):
        """Actualizar barra de progreso"""
        progress_value = (self.question_index / len(self.data)) * 100
        self.progress['value'] = progress_value
    
    def handle_answer(self):
        """Manejar respuesta del usuario"""
        selected_answer = self.answer_var.get()
        
        if not selected_answer:
            messagebox.showwarning("Selección Requerida", 
                                  "Por favor, selecciona una respuesta.")
            return
        
        correct_answer = self.data["Respuesta Correcta"][self.question_index]
        question_text = self.data["Pregunta"][self.question_index]
        
        is_correct = selected_answer == correct_answer
        self.responses.append((question_text, selected_answer, correct_answer, is_correct))
        
        if is_correct:
            self.correct_answers += 1
            messagebox.showinfo("¡Correcto! ✓", "¡Excelente! Respuesta correcta.")
        else:
            messagebox.showerror("Incorrecto ✗", 
                               f"La respuesta correcta era:\n\n{correct_answer}")
        
        self.question_index += 1
        self.show_question()
    
    def skip_question(self):
        """Saltar pregunta actual"""
        if messagebox.askyesno("Saltar Pregunta", "¿Estás seguro de que quieres saltar esta pregunta?"):
            correct_answer = self.data["Respuesta Correcta"][self.question_index]
            question_text = self.data["Pregunta"][self.question_index]
            self.responses.append((question_text, "Sin respuesta", correct_answer, False))
            
            self.question_index += 1
            self.show_question()
    
    def end_game(self):
        """Finalizar el juego"""
        percentage = (self.correct_answers / len(self.data)) * 100
        
        result_msg = f"🎉 Juego Terminado 🎉\n\n"
        result_msg += f"Nombre: {self.student_name}\n"
        result_msg += f"Respuestas correctas: {self.correct_answers} de {len(self.data)}\n"
        result_msg += f"Porcentaje: {percentage:.2f}%\n\n"
        
        if percentage >= 90:
            result_msg += "¡Excelente trabajo! 🌟"
        elif percentage >= 70:
            result_msg += "¡Buen trabajo! 👍"
        elif percentage >= 50:
            result_msg += "Puedes mejorar 📚"
        else:
            result_msg += "Sigue practicando 💪"
        
        messagebox.showinfo("Resultados", result_msg)
        
        self.save_responses_to_word(percentage)
        self.root.destroy()
    
    def save_responses_to_word(self, percentage):
        """Guardar respuestas en documento Word con formato mejorado"""
        try:
            doc = Document()
            
            # Título principal
            title = doc.add_heading('Resultados del Juego de Trivia', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información del estudiante
            doc.add_heading('Información del Estudiante', level=1)
            p = doc.add_paragraph()
            p.add_run('Nombre: ').bold = True
            p.add_run(self.student_name)
            
            p = doc.add_paragraph()
            p.add_run('Fecha: ').bold = True
            p.add_run(datetime.now().strftime("%d/%m/%Y %H:%M"))
            
            p = doc.add_paragraph()
            p.add_run('Archivo: ').bold = True
            p.add_run(os.path.basename(self.file_path))
            
            # Resultados
            doc.add_heading('Resultados Generales', level=1)
            p = doc.add_paragraph()
            p.add_run('Total de preguntas: ').bold = True
            p.add_run(str(len(self.data)))
            
            p = doc.add_paragraph()
            p.add_run('Respuestas correctas: ').bold = True
            p.add_run(f"{self.correct_answers}")
            
            p = doc.add_paragraph()
            p.add_run('Respuestas incorrectas: ').bold = True
            p.add_run(f"{len(self.data) - self.correct_answers}")
            
            p = doc.add_paragraph()
            p.add_run('Porcentaje de aciertos: ').bold = True
            run = p.add_run(f"{percentage:.2f}%")
            run.font.size = Pt(14)
            if percentage >= 70:
                run.font.color.rgb = RGBColor(76, 175, 80)  # Verde
            else:
                run.font.color.rgb = RGBColor(244, 67, 54)  # Rojo
            
            # Detalle de respuestas
            doc.add_page_break()
            doc.add_heading('Detalle de Respuestas', level=1)
            
            for idx, (question, student_answer, correct_answer, is_correct) in enumerate(self.responses, 1):
                doc.add_heading(f'Pregunta {idx}', level=2)
                doc.add_paragraph(question)
                
                p = doc.add_paragraph()
                p.add_run('Tu respuesta: ').bold = True
                run = p.add_run(student_answer)
                if is_correct:
                    run.font.color.rgb = RGBColor(76, 175, 80)
                else:
                    run.font.color.rgb = RGBColor(244, 67, 54)
                
                if not is_correct:
                    p = doc.add_paragraph()
                    p.add_run('Respuesta correcta: ').bold = True
                    run = p.add_run(correct_answer)
                    run.font.color.rgb = RGBColor(76, 175, 80)
                
                p = doc.add_paragraph()
                p.add_run('Estado: ').bold = True
                run = p.add_run('✓ Correcto' if is_correct else '✗ Incorrecto')
                run.font.color.rgb = RGBColor(76, 175, 80) if is_correct else RGBColor(244, 67, 54)
                
                doc.add_paragraph()  # Espacio
            
            # Guardar con nombre único
            filename = f'resultados_{self.student_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            doc.save(filename)
            
            messagebox.showinfo("Guardado", f"Resultados guardados en:\n{filename}")
        except Exception as e:
            messagebox.showerror("Error", f"Error al guardar el documento: {str(e)}")

# Ejecutar aplicación
if __name__ == "__main__":
    root = tk.Tk()
    app = TriviaGame(root)
    root.mainloop()