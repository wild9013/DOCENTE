import pandas as pd
from docx import Document

# Cambia la ruta al archivo Excel a la ubicación correcta
archivo_excel = 'C:/Users/swild/Desktop/python/datos.xlsx'

try:
    # Leer el archivo Excel
    datos = pd.read_excel(archivo_excel)

    # Mostrar las primeras filas del dataframe
    print("Datos leídos:")
    print(datos.head())

    # Asumimos que las columnas de notas son las siguientes
    columnas_notas = ['asitencia1', 'tans1', 'bonificacion1', 'taller1', 
                      'comportamiento1', 'autoevaluacion1', 'examen1']

    # Calcular el promedio por columna
    promedios_por_columna = datos[columnas_notas].mean()

    # Suponemos que la nota mínima para aprobar es 10 (ajusta según tus criterios)
    nota_minima = 10
    estudiantes_aprobados = datos[datos['definitiva'] >= nota_minima]
    estudiantes_reprobados = datos[datos['definitiva'] < nota_minima]

    # Contar la cantidad de estudiantes aprobados y reprobados
    cantidad_aprobados = estudiantes_aprobados.shape[0]
    cantidad_reprobados = estudiantes_reprobados.shape[0]

    # Mostrar resultados
    print(f"\nPromedio general por columna:")
    print(promedios_por_columna)
    print(f"\nCantidad de estudiantes que aprobaron la materia: {cantidad_aprobados}")
    print(f"Cantidad de estudiantes que reprobaron la materia: {cantidad_reprobados}")

    # Crear el informe en Word
    doc = Document()
    doc.add_heading('Informe de Notas de Estudiantes', level=1)

    # Agregar promedios por columna al documento
    doc.add_heading('Promedios por columna:', level=2)
    for columna, promedio in promedios_por_columna.items():
        doc.add_paragraph(f'{columna}: {promedio:.2f}')

    # Agregar cantidad de estudiantes aprobados y reprobados
    doc.add_heading('Cantidad de Estudiantes:', level=2)
    doc.add_paragraph(f'Cantidad de estudiantes que aprobaron la materia: {cantidad_aprobados}')
    doc.add_paragraph(f'Cantidad de estudiantes que reprobaron la materia: {cantidad_reprobados}')

    # Si hay estudiantes reprobados, listar todos los aprendizajes como perdidos
    if cantidad_reprobados > 0:
        aprendizajes_perdidos = datos['aprendizaje'].unique()
        doc.add_heading('Aprendizajes perdidos por estudiantes reprobados:', level=2)
        for aprendizaje in aprendizajes_perdidos:
            doc.add_paragraph(aprendizaje)

    # Guardar el documento
    doc_path = 'C:/Users/swild/Desktop/python/informe_notas.docx'
    doc.save(doc_path)
    print(f"\nInforme guardado en: {doc_path}")

except FileNotFoundError:
    print(f"Error: El archivo no se encontró en la ruta especificada: {archivo_excel}")
except Exception as e:
    print(f"Ocurrió un error: {e}")
