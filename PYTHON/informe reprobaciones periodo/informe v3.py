import pandas as pd
from docx import Document
import os
import tkinter as tk
from tkinter import filedialog, messagebox

def buscar_archivo(titulo_ventana, tipo_archivo_desc, extension):
    """Función para abrir archivos (Lectura)"""
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    
    archivo = filedialog.askopenfilename(
        title=titulo_ventana,
        filetypes=[(tipo_archivo_desc, extension)]
    )
    return archivo

def guardar_archivo(titulo_ventana, carpeta_inicial):
    """Función para guardar archivos (Escritura)"""
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    
    archivo = filedialog.asksaveasfilename(
        title=titulo_ventana,
        initialdir=carpeta_inicial, # Abre por defecto donde estaba el Excel
        defaultextension=".docx",
        filetypes=[("Documento Word", "*.docx")]
    )
    return archivo

# --- CONFIGURACIÓN ---
nota_maxima = 5.0
nota_minima_aprobacion = 3.0 

def main():
    try:
        # ---------------------------------------------------------
        # PASO 1: SELECCIONAR EL EXCEL (DATOS)
        # ---------------------------------------------------------
        print("1. Seleccionando Excel...")
        archivo_excel = buscar_archivo(
            "Paso 1: Selecciona el archivo Excel de NOTAS", 
            "Archivos Excel", 
            "*.xlsx;*.xls"
        )
        
        if not archivo_excel:
            return # Cancelado

        # ---------------------------------------------------------
        # PASO 2: SELECCIONAR LA PLANTILLA (WORD)
        # ---------------------------------------------------------
        print("2. Seleccionando Plantilla...")
        plantilla_word = buscar_archivo(
            "Paso 2: Selecciona la PLANTILLA de informe (Word)", 
            "Documentos Word", 
            "*.docx"
        )

        if not plantilla_word:
            return # Cancelado

        # ---------------------------------------------------------
        # PASO 3: PROCESAMIENTO DE DATOS
        # ---------------------------------------------------------
        print("3. Procesando datos...")
        
        # Leer Excel
        datos = pd.read_excel(archivo_excel, decimal=',') 
        
        # Columnas esperadas
        columnas_notas = ['asitencia1', 'tans1', 'bonificacion1', 'taller1', 
                          'comportamiento1', 'autoevaluacion1', 'examen1']

        # Verificar columnas
        columnas_existentes = [col for col in columnas_notas if col in datos.columns]
        
        if not columnas_existentes:
            messagebox.showerror("Error", "El Excel no tiene las columnas de notas correctas.")
            return

        # Cálculos
        promedios_por_columna = datos[columnas_existentes].mean()
        promedios_por_columna_porcentaje = (promedios_por_columna / nota_maxima) * 100

        if 'definitiva' not in datos.columns:
            messagebox.showerror("Error", "Falta la columna 'definitiva' en el Excel.")
            return

        estudiantes_aprobados = datos[datos['definitiva'] >= nota_minima_aprobacion]
        estudiantes_reprobados = datos[datos['definitiva'] < nota_minima_aprobacion]

        cantidad_aprobados = len(estudiantes_aprobados)
        cantidad_reprobados = len(estudiantes_reprobados)
        total_estudiantes = cantidad_aprobados + cantidad_reprobados
        
        aprobacion_porcentaje = (cantidad_aprobados / total_estudiantes * 100) if total_estudiantes > 0 else 0

        # Textos
        promedios_texto = "\n".join([f"- {col.capitalize()}: {prom:.1f}%" for col, prom in promedios_por_columna_porcentaje.items()])

        if cantidad_reprobados > 0 and 'aprendizaje' in datos.columns:
            aprendizajes_raw = estudiantes_reprobados['aprendizaje'].dropna().unique()
            aprendizajes_texto = "\n".join([f"- {ap}" for ap in aprendizajes_raw])
        else:
            aprendizajes_texto = "No se reportan aprendizajes pendientes."

        # Diccionario
        marcadores = {
            '{{promedios}}': promedios_texto,
            '{{cantidad_aprobados}}': cantidad_aprobados,
            '{{cantidad_reprobados}}': cantidad_reprobados,
            '{{aprendizajes}}': aprendizajes_texto,
            '{{aprobacionporcentaja}}': f"{aprobacion_porcentaje:.1f}%"
        }

        # Generar Word en memoria
        doc = Document(plantilla_word)

        def procesar_reemplazo(parrafo, marcadores):
            for k, v in marcadores.items():
                if k in parrafo.text:
                    parrafo.text = parrafo.text.replace(k, str(v))

        for paragraph in doc.paragraphs:
            procesar_reemplazo(paragraph, marcadores)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        procesar_reemplazo(paragraph, marcadores)

        # ---------------------------------------------------------
        # PASO 4: GUARDAR EL ARCHIVO (USUARIO ELIGE DÓNDE)
        # ---------------------------------------------------------
        print("4. Guardando archivo...")
        
        # Definimos la carpeta inicial como la misma del Excel para comodidad
        carpeta_inicial = os.path.dirname(archivo_excel)
        
        ruta_salida = guardar_archivo("Paso 3: Guardar informe generado como...", carpeta_inicial)

        if not ruta_salida:
            print("Guardado cancelado por el usuario.")
            return

        doc.save(ruta_salida)
        
        print(f"Informe guardado en: {ruta_salida}")
        messagebox.showinfo("¡Éxito!", f"El informe se guardó correctamente en:\n\n{ruta_salida}")

    except PermissionError:
        messagebox.showerror("Error de Permisos", "No se pudo guardar el archivo.\n\nAsegúrate de que no tengas abierto un archivo con el mismo nombre y vuelve a intentarlo.")
    except Exception as e:
        print(f"Error: {e}")
        messagebox.showerror("Error Inesperado", f"Ocurrió un error:\n{e}")

if __name__ == "__main__":
    main()